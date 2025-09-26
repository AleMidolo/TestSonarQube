from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path

    def read_text(self):
        return "\n".join(self._extract_paragraphs_text())

    def write_text(self, content, font_size=12, alignment='left'):
        try:
            doc = Document()
            self._add_paragraph_with_formatting(doc, content, font_size, alignment)
            doc.save(self.file_path)
            return True
        except Exception:
            return False

    def add_heading(self, heading, level=1):
        return self._modify_document(lambda doc: doc.add_heading(heading, level))

    def add_table(self, data):
        return self._modify_document(lambda doc: self._add_table_to_document(doc, data))

    def _extract_paragraphs_text(self):
        doc = Document(self.file_path)
        return [paragraph.text for paragraph in doc.paragraphs]

    def _add_paragraph_with_formatting(self, doc, content, font_size, alignment):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(content)
        run.font.size = Pt(font_size)
        paragraph.alignment = self._get_alignment_value(alignment)

    def _add_table_to_document(self, doc, data):
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        for i, row in enumerate(data):
            for j, cell_value in enumerate(row):
                table.cell(i, j).text = str(cell_value)

    def _modify_document(self, modify_function):
        try:
            doc = Document(self.file_path)
            modify_function(doc)
            doc.save(self.file_path)
            return True
        except Exception:
            return False

    def _get_alignment_value(self, alignment):
        alignment_options = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_options.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)